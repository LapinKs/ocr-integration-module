"""
Конвертер дерева в DOCX (Microsoft Word)
Поддерживает текст, формулы, таблицы
"""

import io
from pathlib import Path
from typing import Dict, Any, List
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib as mpl
mpl.use('Agg')
mpl.rcParams['mathtext.fontset'] = 'stix'
import matplotlib.pyplot as plt
import numpy as np
from PIL import Image
from io import BytesIO

# Константы
DRAW_FORMULA_PLACEHOLDER = False
DEFAULT_FONT_SIZE = 11
DEFAULT_FONT_NAME = "Times New Roman"


class TreeDOCXRenderer:
    """
    Рендерер дерева в DOCX (Microsoft Word)
    """
    
    def __init__(self, font_name: str = DEFAULT_FONT_NAME, font_size: int = DEFAULT_FONT_SIZE):
        self.font_name = font_name
        self.font_size = font_size
        self.document = None
    
    def render(self, page_node, page_width: int, page_height: int) -> bytes:
        """
        Рендеринг страницы в DOCX
        
        Args:
            page_node: корневой узел дерева (тип RIL_PAGE)
            page_width: ширина страницы в пикселях (не используется в Word)
            page_height: высота страницы в пикселях (не используется в Word)
        
        Returns:
            bytes: содержимое DOCX файла
        """
        # Создаём новый документ
        self.document = Document()
        
        # Настройка стилей
        self._setup_styles()
        
        # Рендерим содержимое
        self._render_node(page_node)
        
        # Сохраняем в буфер
        buffer = io.BytesIO()
        self.document.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _setup_styles(self):
        """Настройка стилей документа"""
        # Базовый стиль для абзацев
        style = self.document.styles['Normal']
        style.font.name = self.font_name
        style.font.size = Pt(self.font_size)
        
        # Стиль для формул
        if 'Formula' not in self.document.styles:
            formula_style = self.document.styles.add_style('Formula', 1)  # 1 = WD_STYLE_TYPE_PARAGRAPH
            formula_style.font.name = self.font_name
            formula_style.font.size = Pt(self.font_size + 2)
            formula_style.font.italic = True
            formula_style.paragraph_format.space_before = Pt(6)
            formula_style.paragraph_format.space_after = Pt(6)
            formula_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    def _render_node(self, node):
        """Рекурсивный рендеринг узлов"""
        node_type = node.type
        
        if node_type == "RIL_WORD":
            self._render_word(node)
        
        elif node_type == "RIL_TEXTLINE":
            self._render_textline(node)
        
        elif node_type == "RIL_FORMULA":
            self._render_formula(node)
        
        elif node_type in ["RIL_PAGE", "RIL_TEXT", "RIL_BLOCK", "RIL_LIST_ITEM",
                          "RIL_SECTION_HEADER", "RIL_PAGE_HEADER", "RIL_PAGE_FOOTER"]:
            # Сортируем детей по Y-координате
            sorted_children = sorted(node.children, key=lambda n: n.bbox.y1)
            for child in sorted_children:
                self._render_node(child)
        
        elif node_type == "RIL_TABLE":
            self._render_table(node)
        
        elif node_type == "RIL_TABLE_CELL":
            self._render_table_cell(node)
        
        else:
            for child in node.children:
                self._render_node(child)
    
    def _render_word(self, word_node):
        """Рендеринг слова (добавляет текст в текущий абзац)"""
        text = word_node.data.get("#text", "")
        if not text:
            return
        
        # Добавляем текст в последний абзац или создаём новый
        if len(self.document.paragraphs) == 0:
            self.document.add_paragraph()
        
        paragraph = self.document.paragraphs[-1]
        run = paragraph.add_run(text)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
    
    def _render_textline(self, line_node):
        """Рендеринг строки текста (создаёт новый абзац)"""
        if not line_node.children:
            return
        
        # Сортируем слова по X-координате
        words = sorted([w for w in line_node.children if w.type == "RIL_WORD"],
                      key=lambda w: w.bbox.x1)
        
        if not words:
            return
        
        # Создаём новый абзац
        paragraph = self.document.add_paragraph()
        
        # Группируем слова в текст
        text_parts = []
        for word in words:
            word_text = word.data.get("#text", "")
            if word_text:
                text_parts.append(word_text)
        
        full_text = " ".join(text_parts)
        run = paragraph.add_run(full_text)
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size)
        
        # Выравнивание (по умолчанию - по левому краю)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    def _render_formula(self, formula_node):
        """Рендеринг формулы с использованием matplotlib"""
        latex = formula_node.data.get('latex') or formula_node.data.get('LaTeX')
        
        if latex:
            try:
                # Генерируем изображение формулы
                img_bytes = self._latex_to_image(latex)
                
                # Добавляем изображение в документ
                paragraph = self.document.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                
                # Вставляем изображение
                from docx.shared import Inches
                stream = io.BytesIO(img_bytes)
                inline_shape = run.add_picture(stream, width=Inches(3))
                
            except Exception as e:
                print(f"[DOCX] Formula render error: {e}")
                if DRAW_FORMULA_PLACEHOLDER:
                    paragraph = self.document.add_paragraph()
                    run = paragraph.add_run(f"[Formula: {latex[:50]}...]")
                    run.font.italic = True
                    run.font.color.rgb = None  # серый цвет
        else:
            if DRAW_FORMULA_PLACEHOLDER:
                paragraph = self.document.add_paragraph()
                run = paragraph.add_run("[Formula]")
                run.font.italic = True
    
    def _latex_to_image(self, latex: str, dpi: int = 150) -> bytes:
        """Конвертирует LaTeX в PNG изображение"""
        fig = plt.figure(figsize=(0.01, 0.01))
        fig.patch.set_alpha(0)
        
        # Оборачиваем в $...$ если нужно
        if not (latex.startswith("$") or latex.startswith(r"\[")):
            if any(c in latex for c in ['\\int', '\\sum', '\\frac', '\\begin', '\\lim']):
                latex = r"\[ " + latex + r" \]"
            else:
                latex = r"$" + latex + r"$"
        
        text = fig.text(0, 0, latex, fontsize=14, color='black')
        
        fig.canvas.draw()
        bbox = text.get_window_extent()
        
        width, height = bbox.width / dpi, bbox.height / dpi
        fig.set_size_inches(width + 0.1, height + 0.1)
        
        buf = BytesIO()
        plt.axis("off")
        plt.savefig(buf, dpi=dpi, bbox_inches="tight", pad_inches=0.02, transparent=True)
        plt.close(fig)
        
        buf.seek(0)
        return buf.read()
    
    def _render_table(self, table_node):
        """Рендеринг таблицы"""
        # Собираем ячейки
        cells = [c for c in table_node.children if c.type == "RIL_TABLE_CELL"]
        
        if not cells:
            return
        
        # Определяем структуру таблицы (группируем по Y)
        rows = {}
        for cell in cells:
            # Группируем по Y-координате (с определённым допуском)
            y_center = (cell.bbox.y1 + cell.bbox.y2) // 2
            found_row = False
            for row_y in rows.keys():
                if abs(row_y - y_center) < 20:  # допуск 20 пикселей
                    rows[row_y].append(cell)
                    found_row = True
                    break
            if not found_row:
                rows[y_center] = [cell]
        
        # Сортируем строки по Y
        sorted_rows = sorted(rows.items())
        
        # Создаём таблицу в Word
        num_rows = len(sorted_rows)
        num_cols = max(len(cells) for cells in rows.values())
        
        word_table = self.document.add_table(rows=num_rows, cols=num_cols)
        word_table.style = 'Table Grid'
        word_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Заполняем таблицу
        for row_idx, (row_y, row_cells) in enumerate(sorted_rows):
            # Сортируем ячейки в строке по X
            row_cells.sort(key=lambda c: c.bbox.x1)
            
            for col_idx, cell_node in enumerate(row_cells):
                if col_idx >= num_cols:
                    continue
                
                cell = word_table.cell(row_idx, col_idx)
                cell.text = ""
                
                # Рендерим содержимое ячейки
                for child in cell_node.children:
                    self._render_node_in_cell(cell, child)
    
    def _render_node_in_cell(self, cell, node):
        """Рендеринг узла внутри ячейки таблицы"""
        node_type = node.type
        
        if node_type == "RIL_WORD":
            text = node.data.get("#text", "")
            if text:
                # Добавляем текст с пробелом
                current_text = cell.text
                if current_text and not current_text.endswith(' '):
                    cell.text = current_text + " " + text
                else:
                    cell.text = current_text + text
        
        elif node_type == "RIL_TEXTLINE":
            # Рекурсивно рендерим слова
            words = sorted([w for w in node.children if w.type == "RIL_WORD"],
                          key=lambda w: w.bbox.x1)
            for word in words:
                self._render_node_in_cell(cell, word)
        
        elif node_type == "RIL_FORMULA":
            latex = node.data.get('latex') or node.data.get('LaTeX')
            if latex:
                try:
                    # Для формул в таблицах пока используем текстовое представление
                    # (можно добавить изображения, но это сложнее)
                    cell.text = cell.text + f" [{latex[:30]}...] " if len(latex) > 30 else cell.text + f" {latex} "
                except:
                    cell.text = cell.text + " [Formula] "
        
        else:
            for child in node.children:
                self._render_node_in_cell(cell, child)
    
    def _render_table_cell(self, cell_node):
        """Рендеринг ячейки таблицы (используется в _render_table)"""
        # Основная логика в _render_table
        pass


def render_page_to_docx(page_node, page_width: int, page_height: int) -> bytes:
    """
    Удобная функция для рендеринга страницы в DOCX
    
    Args:
        page_node: корневой узел дерева
        page_width: ширина страницы (не используется)
        page_height: высота страницы (не используется)
    
    Returns:
        bytes: содержимое DOCX файла
    """
    renderer = TreeDOCXRenderer()
    return renderer.render(page_node, page_width, page_height)